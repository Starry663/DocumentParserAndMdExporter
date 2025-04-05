# word_parser.py
from docx import Document
import zipfile, os
from PIL import Image

def parse_docx(docx_path, output_dir):
    # 确保输出目录存在，如果没有则创建
    img_dir = os.path.join(output_dir, "images")
    os.makedirs(img_dir, exist_ok=True)

    # 提取 Word 文档中的所有图片文件
    with zipfile.ZipFile(docx_path, 'r') as archive:
        # 筛选出存储图片的文件，如 "word/media/image1.png"
        for file in archive.namelist():
            if file.startswith("word/media/"):
                # 将图片提取到输出 images 目录
                archive.extract(file, img_dir)
                # zip 提取会保留原始路径，我们需要将图片移到 images 根目录并改名
                extracted_path = os.path.join(img_dir, file)
                img_name = os.path.basename(file)  # 原始图片文件名
                new_path = os.path.join(img_dir, os.path.basename(docx_path) + "_" + img_name)
                os.replace(extracted_path, new_path)
                # 可选：将图片统一转换为PNG格式
                try:
                    im = Image.open(new_path)
                    png_path = os.path.splitext(new_path)[0] + ".png"
                    im.save(png_path, format="PNG")
                    os.remove(new_path)  # 删除原格式文件
                    new_path = png_path
                except Exception as e:
                    print(f"图片转换PNG格式失败: {e}")
                # 到这里，图片已保存为 new_path 路径（PNG 格式）

    # 使用 python-docx 读取文档文本和表格
    doc = Document(docx_path)
    text_runs = []    # 存储所有段落文本的列表
    for para in doc.paragraphs:
        if para.text:  # 去除空段落
            text_runs.append(para.text)

    tables_data = []  # 存储所有表格的数据
    for table in doc.tables:
        # 提取表格的每一行，每一行又是一个单元格列表
        table_rows = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            table_rows.append(cells_text)
        tables_data.append(table_rows)

    # 获取刚才保存的所有图片文件路径列表
    image_files = []
    for fname in os.listdir(img_dir):
        # 只添加当前docx相关且为图片格式的文件
        if os.path.basename(docx_path) in fname:
            image_files.append(os.path.join("images", fname))
    # 返回提取的文本、表格和图片路径
    return text_runs, tables_data, image_files
