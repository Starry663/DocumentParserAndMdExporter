import sys, os
import fitz                  # PyMuPDF for PDF parsing
import docx                  # python-docx for Word parsing
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QSplitter, QTreeWidget, QTreeWidgetItem,
    QPlainTextEdit, QTextEdit, QFileDialog
)
from PyQt5.QtCore import Qt, QUrl

def parse_document(file_path):
    """
    Parse the given Word (.docx) or PDF document and return:
      - structure: list of (level, title) representing the document's outline (headings or bookmarks).
      - full_text: the complete plain text content of the document (no images).
      - md_path: file path to a generated Markdown file with the document's content and image links.
    """
    structure = []
    full_text = ''
    md_path = ''
    if file_path.lower().endswith(('.docx', '.doc')):
        # --- Parse Word document ---
        doc = docx.Document(file_path)
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        # Extract outline structure (headings) and full text content
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Identify heading level by style name (supports English "Heading X" or Chinese "标题 X")
            style_name = para.style.name if para.style else ''
            level = None
            for lvl in range(1, 10):
                if style_name.lower().startswith(f'heading {lvl}') or style_name == f'Heading {lvl}' or style_name == f'标题 {lvl}':
                    level = lvl
                    break
            if level:  # It's a heading, add to structure
                structure.append((level, text))
            # Append text to full content (with newline)
            full_text += text + '\n'
        # Prepare directory for extracted images
        img_dir = os.path.join(os.path.dirname(file_path), f"{base_name}_images")
        os.makedirs(img_dir, exist_ok=True)
        md_lines = []
        # Build Markdown text content
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ''
            if text:
                # Determine if paragraph is a heading and format accordingly
                level = None
                for lvl in range(1, 10):
                    if style_name.lower().startswith(f'heading {lvl}') or style_name == f'Heading {lvl}' or style_name == f'标题 {lvl}':
                        level = lvl
                        break
                if level:
                    md_lines.append('#' * level + ' ' + text)   # Convert heading to Markdown syntax
                else:
                    md_lines.append(text)
            else:
                # Empty paragraph (could be a blank line or an image placeholder)
                md_lines.append('')
        # Extract images from the Word document and save them, adding references in Markdown
        image_counter = 1
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:  # Identify image relationships in the docx
                image_bytes = rel.target_part.blob
                # Determine image file extension
                ext = os.path.splitext(rel.target_ref)[1]
                if not ext:
                    ext = '.png'
                img_filename = f"image_{image_counter}{ext}"
                img_path = os.path.join(img_dir, img_filename)
                # Save image to the images directory
                with open(img_path, 'wb') as img_file:
                    img_file.write(image_bytes)
                # Add an image reference to the Markdown content (relative path)
                md_lines.append(f"![image_{image_counter}]({os.path.basename(img_dir)}/{img_filename})")
                image_counter += 1
        # Write the Markdown content to a .md file
        md_content = '\n\n'.join(md_lines)
        md_path = os.path.join(os.path.dirname(file_path), f"{base_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        return structure, full_text, md_path

    elif file_path.lower().endswith('.pdf'):
        # --- Parse PDF document ---
        doc = fitz.open(file_path)
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        # Extract outline structure from PDF bookmarks (table of contents)
        toc = doc.get_toc(simple=True)  # list of [level, title, page] entries
        structure = [(lvl, title) for (lvl, title, page) in toc]
        # Extract full text from all pages
        for page in doc:
            full_text += page.get_text("text")
        # Prepare directory for extracted images
        img_dir = os.path.join(os.path.dirname(file_path), f"{base_name}_images")
        os.makedirs(img_dir, exist_ok=True)
        md_lines = []
        # Build Markdown text content from PDF text, preserving paragraph breaks
        for line in full_text.splitlines():
            if line.strip() == "":
                md_lines.append("")  # blank line for paragraph separation
            else:
                md_lines.append(line)
        # Extract images from each page and save them, adding references in Markdown
        image_counter = 1
        for page_number in range(len(doc)):
            page = doc[page_number]
            for img in page.get_images(full=True):  # list of images on this page
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                ext = base_image.get("ext", "png")
                img_filename = f"image_{page_number+1}_{image_counter}.{ext}"
                img_path = os.path.join(img_dir, img_filename)
                # Save the image file
                with open(img_path, "wb") as img_file:
                    img_file.write(image_bytes)
                # Add image reference to Markdown (relative path)
                md_lines.append(f"![image_{page_number+1}_{image_counter}]({os.path.basename(img_dir)}/{img_filename})")
                image_counter += 1
        # Write the Markdown content to a .md file
        md_content = '\n\n'.join(md_lines)
        md_path = os.path.join(os.path.dirname(file_path), f"{base_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        return structure, full_text, md_path

    else:
        raise ValueError("Unsupported file format: must be .docx or .pdf")


class MainWindow(QMainWindow):
    def __init__(self, file_path):
        super().__init__()
        self.setWindowTitle("文档查看")  # Main window title

        # Parse the selected document to get outline, text, and markdown content
        structure, full_text, md_path = parse_document(file_path)

        # Create a splitter to divide the window into three panels
        splitter = QSplitter(Qt.Horizontal)

        # Left Panel: QTreeWidget for the document outline (headings/bookmarks)
        tree_widget = QTreeWidget()
        tree_widget.setHeaderHidden(True)  # no header needed for outline
        # Populate the tree with the outline structure
        parents = {0: None}
        for level, title in structure:
            item = QTreeWidgetItem([title])
            parent_item = parents.get(level - 1)
            if parent_item is None:
                # This is a top-level item
                tree_widget.addTopLevelItem(item)
            else:
                # Add as a child to the appropriate parent item
                parent_item.addChild(item)
            parents[level] = item  # update the current level parent
        tree_widget.expandAll()  # expand all nodes to show the full hierarchy

        # Middle Panel: QPlainTextEdit for full text content
        text_edit = QPlainTextEdit()
        text_edit.setPlainText(full_text)
        text_edit.setReadOnly(True)  # make it view-only

        # Right Panel: QTextEdit for Markdown content (with images preview)
        markdown_edit = QTextEdit()
        markdown_edit.setReadOnly(True)
        # Load the generated Markdown file content
        md_content = ""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
        except Exception as e:
            md_content = f"无法加载 Markdown 文件: {e}"
        # Set base URL for image references so that local images can be displayed
        base_url = QUrl.fromLocalFile(os.path.join(os.path.dirname(md_path), ''))
        markdown_edit.document().setBaseUrl(base_url)
        # Display the Markdown content (Qt will render formatting and images)
        markdown_edit.setMarkdown(md_content)

        # Add all three panels to the splitter
        splitter.addWidget(tree_widget)
        splitter.addWidget(text_edit)
        splitter.addWidget(markdown_edit)
        # Adjust initial proportions of the splitter (optional)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 2)
        splitter.setStretchFactor(2, 2)

        # Set the splitter as the central widget of the main window
        self.setCentralWidget(splitter)


# If running this module directly, allow selecting a file to open
if __name__ == "__main__":
    app = QApplication(sys.argv)
    # If a file path is provided as a command-line argument, use it; otherwise, open a file dialog
    file_path = sys.argv[1] if len(sys.argv) > 1 else None
    if not file_path:
        # Only allow selecting Word or PDF files
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            None, "选择文档", "", "Word/PDF 文件 (*.docx *.pdf)", options=options
        )
    if not file_path:
        sys.exit("未选择文件。")
    # Show the main window with the selected document
    window = MainWindow(file_path)
    window.show()
    sys.exit(app.exec_())

