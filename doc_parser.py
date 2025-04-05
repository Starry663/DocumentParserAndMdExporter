import os
import sys
import re
import fitz            # PyMuPDF，用于 PDF 文本和图片解析
import pdfplumber      # 用于 PDF 表格提取
from docx import Document
from docx.parts.image import ImagePart

def get_output_dir():
    """获取项目主目录下的 output 输出文件夹路径，如不存在则创建。"""
    if getattr(sys, 'frozen', False):  # 检测是否为 PyInstaller 打包后的运行环境
        exe_dir = os.path.dirname(sys.executable)       # exe 所在目录
        parent_dir = os.path.dirname(exe_dir)           # 上一级目录
        # 判断 exe 是否在 dist 或其子目录下，以定位项目主目录
        if os.path.basename(exe_dir).lower() == 'dist':
            # exe 位于项目/dist 目录下（如 onefile 模式）
            project_dir = parent_dir
        elif os.path.basename(parent_dir).lower() == 'dist':
            # exe 位于项目/dist/子目录 下（如 onefolder 模式）
            project_dir = os.path.dirname(parent_dir)
        else:
            # exe 不在预期的 dist 结构内（可能被移动），则退而求其次取其上级目录
            project_dir = parent_dir
    else:
        # 源码运行，直接取当前文件所在目录作为项目主目录
        project_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(project_dir, "output")
    os.makedirs(output_path, exist_ok=True)  # 确保 output 文件夹存在
    return output_path

def doc_to_markdown(doc_path):
    """解析 Word 文档并返回对应的 Markdown 文本字符串。"""
    doc = Document(doc_path)
    output_dir = get_output_dir()
    base_name = os.path.splitext(os.path.basename(doc_path))[0]  # 输入文件无扩展名部分
    md_lines = []

    # 提取段落文本
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            md_lines.append(text)

    # 提取表格并转换为 Markdown 表格格式
    for table in doc.tables:
        rows = []
        for row in table.rows:
            # 获取该行所有单元格文本，去除首尾空白
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            # 将表格第一个行作为表头
            header = rows[0]
            col_count = len(header)
            md_lines.append('|' + '|'.join(header) + '|')
            md_lines.append('|' + '|'.join(['---'] * col_count) + '|')
            for row in rows[1:]:
                md_line = '|' + '|'.join(row) + '|'
                md_lines.append(md_line)

    # 提取图片并保存到 output 文件夹，Markdown 中插入对应引用
    image_count = 0
    for rel in doc.part.rels.values():
        if isinstance(rel._target, ImagePart):
            image_count += 1
            image_bytes = rel._target.blob
            # 使用原始图片文件的扩展名（确保保存格式正确）
            img_ext = os.path.splitext(rel._target.filename)[1]  # 包含点号，例如 ".png" 或 ".jpg"
            # 构造输出图片文件名：{文档名}_image{序号}.{扩展名}
            image_filename = f"{base_name}_image{image_count}{img_ext}"
            image_path = os.path.join(output_dir, image_filename)
            with open(image_path, "wb") as img_file:
                img_file.write(image_bytes)
            # 在 Markdown 内容中添加图片引用（图片存放在与 Markdown 同一目录下）
            md_lines.append(f"![{image_filename}]({image_filename})")

    # 组合所有行，段落之间空行分隔，形成最终 Markdown 文本
    markdown_text = "\n\n".join(md_lines)
    return markdown_text

def pdf_to_markdown(pdf_path):
    """解析 PDF 文档并返回对应的 Markdown 文本字符串。"""
    output_dir = get_output_dir()
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    md_lines = []

    # 打开 PDF（PyMuPDF 用于文本和图片提取，pdfplumber 用于表格提取）
    fitz_doc = fitz.open(pdf_path)
    pdf = pdfplumber.open(pdf_path)

    for page_index in range(len(fitz_doc)):
        # 提取页面文本内容
        page = fitz_doc[page_index]
        text = page.get_text().strip()
        if text:
            md_lines.append(text)
        # 提取页面中的表格，并转换为 Markdown 表格
        tables = pdf.pages[page_index].extract_tables()
        for table in tables:
            if table:
                # 将表格第一行作为表头
                header = [(cell if cell is not None else "").strip() for cell in table[0]]
                col_count = len(header)
                md_lines.append('|' + '|'.join(header) + '|')
                md_lines.append('|' + '|'.join(['---'] * col_count) + '|')
                for row in table[1:]:
                    row_cells = [(cell if cell is not None else "").strip() for cell in row]
                    md_lines.append('|' + '|'.join(row_cells) + '|')
        # 提取页面中的图片
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]  # 获取图像xref
            try:
                base_image = fitz_doc.extract_image(xref)
            except Exception:
                base_image = None
            if base_image:
                image_bytes = base_image.get("image")
                img_ext = base_image.get("ext", "png")  # 图片扩展名，如 'png', 'jpg' 等
                # 构造输出图片文件名：{PDF名}_image{序号}.{扩展名}
                # 使用全局计数避免不同页面图片重名
                image_index = len([name for name in os.listdir(output_dir) if name.startswith(f"{base_name}_image")]) + 1
                image_filename = f"{base_name}_image{image_index}.{img_ext}"
                image_path = os.path.join(output_dir, image_filename)
                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)
                md_lines.append(f"![{image_filename}]({image_filename})")

    # 关闭 PDF 文件
    fitz_doc.close()
    pdf.close()
    # 组合 Markdown 内容
    markdown_text = "\n\n".join(md_lines)
    return markdown_text

def parse_document(file_path):
    """解析给定的文件（Word 或 PDF），生成 Markdown 文件并返回其路径。"""
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext in [".doc", ".docx"]:
        md_content = doc_to_markdown(file_path)
    elif file_ext == ".pdf":
        md_content = pdf_to_markdown(file_path)
    else:
        raise ValueError("不支持的文件格式！")

    # 将 Markdown 内容写入文件，文件名与输入文档同名（扩展名为.md）
    output_dir = get_output_dir()
    md_filename = os.path.splitext(os.path.basename(file_path))[0] + ".md"
    md_path = os.path.join(output_dir, md_filename)
    with open(md_path, "w", encoding="utf-8") as md_file:
        md_file.write(md_content)
    return md_path



